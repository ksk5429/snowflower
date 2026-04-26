"""generate_report.py — produces snowflower_strategy_report.docx (~25 pages, plain language).

Demonstrates advanced python-docx use plus deliberate plain-language writing:
short sentences, simple words, defined acronyms, concrete examples, "in plain
words" callouts, glossary, FAQ section.

Usage:
    python generate_report.py
"""

from __future__ import annotations

import sys
from datetime import datetime
from pathlib import Path

# Force UTF-8 stdout for Windows codepage compatibility (Korean strings).
for _stream in (sys.stdout, sys.stderr):
    try:
        _stream.reconfigure(encoding="utf-8", errors="replace")  # type: ignore[attr-defined]
    except (AttributeError, ValueError):
        pass

import matplotlib

matplotlib.use("Agg")
import matplotlib.patches as mpatches
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

REPORT_PATH = Path("snowflower_strategy_report.docx")
ASSETS_DIR = Path("_report_assets")
ASSETS_DIR.mkdir(exist_ok=True)

DOC_TITLE = "snowflower"
DOC_SUBTITLE = "An independent channel about hands-free technology"
DOC_VERSION = "Strategy Report v2.0 — plain-language edition"
DOC_DATE = datetime.now().strftime("%Y-%m-%d")
DOC_AUTHOR = "snowflower editorial"

# Color palette
PRIMARY_HEX = "1A3A52"
ACCENT_HEX = "E8F2F9"
MUTED_HEX = "6B7380"
CALLOUT_HEX = "FFF8E7"
DANGER_HEX = "FCE7E7"
SUCCESS_HEX = "E7F5E7"
INFO_HEX = "E7EEF5"

PRIMARY_RGB = RGBColor(0x1A, 0x3A, 0x52)
MUTED_RGB = RGBColor(0x6B, 0x73, 0x80)
TEXT_RGB = RGBColor(0x22, 0x22, 0x22)


# ---------------------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------------------


def _set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_table_borders(table, color: str = "CCCCCC", size: str = "4") -> None:
    tbl = table._tbl
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    tbl_borders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:color"), color)
        tbl_borders.append(border)
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_pr.append(tbl_borders)


def _add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def _add_total_pages_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "NUMPAGES"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def _add_toc_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r'TOC \o "1-3" \h \z \u'
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = (
        "Open this in Microsoft Word, then right-click here and choose "
        "'Update Field' to fill in the table of contents."
    )
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    run._r.append(placeholder)
    run._r.append(fld_char3)


# ---------------------------------------------------------------------------
# Style setup
# ---------------------------------------------------------------------------


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TEXT_RGB
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(20)
    h1.font.bold = True
    h1.font.color.rgb = PRIMARY_RGB
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = PRIMARY_RGB
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(11.5)
    h3.font.bold = True
    h3.font.color.rgb = TEXT_RGB
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(2)
    h3.paragraph_format.keep_with_next = True

    for section in doc.sections:
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)


def set_core_properties(doc: Document) -> None:
    cp = doc.core_properties
    cp.title = f"{DOC_TITLE} — {DOC_VERSION}"
    cp.author = DOC_AUTHOR
    cp.subject = "Independent channel about hands-free assistive technology"
    cp.keywords = (
        "accessibility, assistive technology, MouthPad, augmental, snowflower, "
        "editorial, content engine, plain language"
    )
    cp.comments = "Generated by generate_report.py — plain-language edition"


# ---------------------------------------------------------------------------
# Document-building helpers
# ---------------------------------------------------------------------------


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    if level == 1 and doc.paragraphs and doc.paragraphs[-1].text.strip():
        add_page_break(doc)
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str, *, bold: bool = False, italic: bool = False,
                  align=None, size_pt: float | None = None,
                  color: RGBColor | None = None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if color is not None:
        run.font.color.rgb = color


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]],
              header_fill: str = PRIMARY_HEX, alt_fill: str = "F7F9FC",
              col_widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_borders(table)

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        _set_cell_shading(cell, header_fill)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = table.rows[r].cells[c]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if r % 2 == 0:
                _set_cell_shading(cell, alt_fill)

    if col_widths:
        for r in table.rows:
            for c, w in enumerate(col_widths):
                if c < len(r.cells):
                    r.cells[c].width = Inches(w)

    doc.add_paragraph()


def add_callout(doc: Document, title: str, body: str, fill_hex: str = CALLOUT_HEX) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, fill_hex)
    _set_table_borders(table, color="DDDDDD")

    cell.text = ""
    p1 = cell.paragraphs[0]
    r1 = p1.add_run(title)
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = PRIMARY_RGB

    p2 = cell.add_paragraph()
    r2 = p2.add_run(body)
    r2.font.size = Pt(9.5)
    doc.add_paragraph()


def add_plain_words(doc: Document, body: str) -> None:
    """Special callout that summarizes the prior section in plain language."""
    add_callout(doc, "In plain words", body, INFO_HEX)


def add_figure(doc: Document, image_path: Path, caption: str,
               width_inches: float = 6.0) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    cap_run.italic = True
    cap_run.font.size = Pt(9)
    cap_run.font.color.rgb = MUTED_RGB
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Cover + header / footer
# ---------------------------------------------------------------------------


def build_cover(doc: Document) -> None:
    for _ in range(6):
        doc.add_paragraph()

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run(DOC_TITLE)
    r.bold = True
    r.font.size = Pt(56)
    r.font.color.rgb = PRIMARY_RGB

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p_sub.add_run(DOC_SUBTITLE)
    r2.italic = True
    r2.font.size = Pt(14)
    r2.font.color.rgb = MUTED_RGB

    doc.add_paragraph()
    doc.add_paragraph()

    p_ver = doc.add_paragraph()
    p_ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p_ver.add_run(DOC_VERSION)
    r3.font.size = Pt(13)
    r3.bold = True
    r3.font.color.rgb = PRIMARY_RGB

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p_date.add_run(DOC_DATE)
    r4.font.size = Pt(11)
    r4.font.color.rgb = MUTED_RGB

    for _ in range(12):
        doc.add_paragraph()

    p_author = doc.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r5 = p_author.add_run(
        f"Prepared by {DOC_AUTHOR}\nrepository: github.com/ksk5429/snowflower"
    )
    r5.font.size = Pt(10)
    r5.font.color.rgb = MUTED_RGB


def configure_header_footer(doc: Document) -> None:
    for section in doc.sections[1:]:
        header = section.header
        h_para = header.paragraphs[0]
        h_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        h_run = h_para.add_run(f"{DOC_TITLE} · {DOC_VERSION}")
        h_run.font.size = Pt(9)
        h_run.font.color.rgb = MUTED_RGB
        h_run.italic = True

        footer = section.footer
        f_para = footer.paragraphs[0]
        f_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_pre = f_para.add_run("Page ")
        run_pre.font.size = Pt(9)
        run_pre.font.color.rgb = MUTED_RGB
        _add_page_number_field(f_para)
        run_mid = f_para.add_run(" of ")
        run_mid.font.size = Pt(9)
        run_mid.font.color.rgb = MUTED_RGB
        _add_total_pages_field(f_para)
        run_post = f_para.add_run(f"   ·   {DOC_DATE}")
        run_post.font.size = Pt(9)
        run_post.font.color.rgb = MUTED_RGB


# ---------------------------------------------------------------------------
# Charts
# ---------------------------------------------------------------------------


def chart_platform_matrix() -> Path:
    platforms = [
        ("YouTube", 8, 9, 1),
        ("LinkedIn", 8, 10, 1),
        ("Bluesky", 9, 8, 1),
        ("Beehiiv", 9, 9, 1),
        ("Threads", 7, 5, 2),
        ("Instagram", 6, 7, 2),
        ("TikTok", 5, 7, 2),
        ("X", 6, 4, 2),
        ("Reddit", 3, 6, 3),
        ("Naver Blog", 3, 8, 3),
        ("Stibee", 9, 8, 3),
    ]
    tier_colors = {1: "#1A3A52", 2: "#5B7FA8", 3: "#B0B8C0"}

    fig, ax = plt.subplots(figsize=(8, 5), dpi=140)
    for name, x, y, tier in platforms:
        ax.scatter(x, y, s=220, c=tier_colors[tier], alpha=0.85, edgecolor="white", linewidth=1.5)
        ax.annotate(name, (x, y), xytext=(8, 4), textcoords="offset points", fontsize=9, color="#222")

    ax.set_xlabel("How easy to automate (10 = very easy)", fontsize=10)
    ax.set_ylabel("How well it fits our audience (10 = perfect fit)", fontsize=10)
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 11)
    ax.set_xticks(range(0, 12, 2))
    ax.set_yticks(range(0, 12, 2))
    ax.grid(True, linestyle=":", alpha=0.4)
    ax.set_title("Platform map: where to invest first", fontsize=11, pad=12)

    legend_handles = [
        plt.Line2D([0], [0], marker="o", color="w", markerfacecolor=tier_colors[1],
                   markersize=12, label="Tier 1: do first"),
        plt.Line2D([0], [0], marker="o", color="w", markerfacecolor=tier_colors[2],
                   markersize=12, label="Tier 2: after Tier 1 works"),
        plt.Line2D([0], [0], marker="o", color="w", markerfacecolor=tier_colors[3],
                   markersize=12, label="Tier 3: manual / Korea"),
    ]
    ax.legend(handles=legend_handles, loc="lower right", fontsize=9, framealpha=0.95)

    plt.tight_layout()
    out = ASSETS_DIR / "platform_matrix.png"
    plt.savefig(out, dpi=140, bbox_inches="tight")
    plt.close(fig)
    return out


def chart_budget_breakdown() -> Path:
    items = [
        ("Cloud / GitHub Actions", 6),
        ("Beehiiv (after 2,500 subs)", 39),
        ("ElevenLabs (Korean voice)", 22),
        ("fal.ai (image generation)", 18),
        ("X API Basic", 17),
        ("Captions / misc tools", 8),
    ]
    items_sorted = sorted(items, key=lambda x: x[1])
    labels = [i[0] for i in items_sorted]
    vals = [i[1] for i in items_sorted]

    fig, ax = plt.subplots(figsize=(8, 4.2), dpi=140)
    bars = ax.barh(labels, vals, color="#1A3A52", alpha=0.85, edgecolor="white")
    ax.set_xlabel("US dollars per month", fontsize=10)
    ax.set_xlim(0, max(vals) * 1.25)
    for bar, val in zip(bars, vals):
        ax.text(val + 1, bar.get_y() + bar.get_height() / 2,
                f"${val}", va="center", fontsize=9, color="#222")

    total = sum(vals)
    ax.set_title(f"Where the money goes each month — total about ${total} (limit: $50–100)",
                 fontsize=11, pad=12)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.grid(True, axis="x", linestyle=":", alpha=0.4)

    plt.tight_layout()
    out = ASSETS_DIR / "budget_breakdown.png"
    plt.savefig(out, dpi=140, bbox_inches="tight")
    plt.close(fig)
    return out


def chart_90_day_timeline() -> Path:
    workstreams = [
        ("Set up: domain, gmail, phone", 0, 1, "#5B7FA8"),
        ("Make accounts (Tier 1)", 1, 3, "#5B7FA8"),
        ("Pre-make 4 episodes", 1, 3, "#1A3A52"),
        ("Methodology page + first publish", 3, 5, "#1A3A52"),
        ("Weekly publishing rhythm", 4, 13, "#1A3A52"),
        ("Newsletter every week", 4, 13, "#7A9FBC"),
        ("LinkedIn 2x/week + daily Bluesky", 4, 13, "#7A9FBC"),
        ("CSUN 2026 conference", 9, 10, "#C8451A"),
        ("Edit conference videos (6 episodes)", 10, 12, "#1A3A52"),
        ("Pitch to augmental.tech", 12, 13, "#C8451A"),
    ]

    fig, ax = plt.subplots(figsize=(8.6, 4.6), dpi=140)
    for i, (name, start, end, color) in enumerate(workstreams):
        ax.barh(i, end - start, left=start, color=color, alpha=0.88,
                edgecolor="white", height=0.65)
        ax.text(start + 0.1, i, name, va="center", fontsize=8.5,
                color="white" if (end - start) > 2 else "#222")

    ax.set_yticks(range(len(workstreams)))
    ax.set_yticklabels([f"W{ws[1]}-{ws[2]}" for ws in workstreams], fontsize=8.5, color="#555")
    ax.invert_yaxis()
    ax.set_xlabel("Week", fontsize=10)
    ax.set_xlim(0, 14)
    ax.set_xticks(range(0, 14, 2))
    ax.set_title("90-day plan — conference at week 9, pitch at week 12", fontsize=11, pad=12)
    ax.grid(True, axis="x", linestyle=":", alpha=0.4)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)

    plt.tight_layout()
    out = ASSETS_DIR / "timeline_90day.png"
    plt.savefig(out, dpi=140, bbox_inches="tight")
    plt.close(fig)
    return out


def chart_buyer_journey() -> Path:
    """Text-box flow showing how a MouthPad gets purchased."""
    fig, ax = plt.subplots(figsize=(8.4, 5), dpi=140)
    ax.axis("off")

    boxes = [
        # (text, x, y, color)
        ("User has\nmobility limitation\n(SCI, ALS, MS, CP)", 1, 4.5, "#E8F2F9"),
        ("Sees snowflower video\nor newsletter article", 4, 4.5, "#FFF8E7"),
        ("Talks to OT or SLP\n(occupational or speech\ntherapist)", 7, 4.5, "#E8F2F9"),
        ("Therapist evaluates\nand writes\nLetter of Medical Necessity", 7, 2.5, "#E8F2F9"),
        ("Insurance / VA / school\nreviews paperwork", 4, 2.5, "#FFF8E7"),
        ("Approved → device shipped\nor Denied → appeal", 1, 2.5, "#E7F5E7"),
    ]

    for text, x, y, color in boxes:
        rect = mpatches.FancyBboxPatch(
            (x - 1.3, y - 0.6), 2.6, 1.2,
            boxstyle="round,pad=0.05", linewidth=1.5,
            edgecolor="#1A3A52", facecolor=color
        )
        ax.add_patch(rect)
        ax.text(x, y, text, ha="center", va="center", fontsize=8.5, color="#222")

    # arrows
    arrows = [
        ((2.3, 4.5), (2.7, 4.5)),     # user -> snowflower
        ((5.3, 4.5), (5.7, 4.5)),     # snowflower -> therapist
        ((7, 3.9), (7, 3.1)),         # therapist down to LMN
        ((5.7, 2.5), (5.3, 2.5)),     # LMN -> insurance
        ((2.7, 2.5), (2.3, 2.5)),     # insurance -> device
    ]
    for (sx, sy), (ex, ey) in arrows:
        ax.annotate("", xy=(ex, ey), xytext=(sx, sy),
                    arrowprops=dict(arrowstyle="->", color="#1A3A52", lw=1.8))

    ax.set_xlim(-0.5, 9)
    ax.set_ylim(1.5, 5.5)
    ax.set_title("How a MouthPad actually gets purchased (US path shown)",
                 fontsize=11, pad=12)

    plt.tight_layout()
    out = ASSETS_DIR / "buyer_journey.png"
    plt.savefig(out, dpi=140, bbox_inches="tight")
    plt.close(fig)
    return out


def chart_revenue_mix() -> Path:
    """Stacked area: how revenue mix changes month 1 to month 12."""
    months = list(range(1, 13))
    # Each list = monthly $ from that source
    adsense = [0, 0, 5, 15, 30, 50, 70, 90, 110, 140, 170, 200]
    affiliate = [0, 0, 0, 5, 10, 20, 35, 50, 70, 90, 110, 130]
    newsletter_sponsor = [0, 0, 0, 0, 50, 150, 250, 400, 600, 800, 1000, 1200]
    patreon = [0, 0, 0, 50, 100, 200, 350, 500, 700, 900, 1100, 1300]
    anchor_sponsor = [0, 0, 0, 0, 0, 0, 0, 1000, 2000, 3000, 4000, 5000]

    fig, ax = plt.subplots(figsize=(8.4, 4.6), dpi=140)
    ax.stackplot(
        months,
        adsense, affiliate, newsletter_sponsor, patreon, anchor_sponsor,
        labels=["YouTube AdSense", "Amazon affiliate", "Newsletter sponsors",
                "Patreon membership", "Anchor sponsor (e.g. augmental)"],
        colors=["#B0B8C0", "#7A9FBC", "#5B7FA8", "#1A3A52", "#C8451A"],
        alpha=0.9,
    )
    ax.set_xlabel("Month", fontsize=10)
    ax.set_ylabel("Revenue (US dollars per month)", fontsize=10)
    ax.set_xticks(months)
    ax.set_title("How revenue should grow over the first 12 months (estimate)",
                 fontsize=11, pad=12)
    ax.legend(loc="upper left", fontsize=9, framealpha=0.95)
    ax.grid(True, linestyle=":", alpha=0.4)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)

    plt.tight_layout()
    out = ASSETS_DIR / "revenue_mix.png"
    plt.savefig(out, dpi=140, bbox_inches="tight")
    plt.close(fig)
    return out


# ---------------------------------------------------------------------------
# Sections — plain-language edition
# ---------------------------------------------------------------------------


def section_readers_guide(doc: Document) -> None:
    add_heading(doc, "Reader's Guide", level=1)
    doc.add_paragraph(
        "This report is written in plain English. The goal is that anyone who reads "
        "it — whether a doctor, an investor, a software engineer, or a friend — can "
        "understand the strategy without needing a marketing background."
    )

    add_heading(doc, "How to use this report", level=2)
    add_bullets(
        doc,
        [
            "If you only have 5 minutes: read the Executive Summary (Section 1) and the "
            "Big Picture (Section 2).",
            "If you want to know what we are actually building: Sections 6 (Engine) and "
            "8 (Sample Content) show concrete examples.",
            "If you want to know the plan: Section 12 (90-Day Plan) is a week-by-week list.",
            "If you have questions: jump to Section 16 (FAQ) — it answers the most "
            "common ones.",
            "If a word is unfamiliar: check the Glossary at the front.",
        ],
    )

    add_heading(doc, "Conventions used", level=2)
    add_bullets(
        doc,
        [
            "Yellow boxes labeled 'In plain words' summarize the section above in even "
            "simpler language.",
            "Blue boxes are background information.",
            "Green boxes show what success looks like.",
            "Red boxes show risks or things to watch out for.",
            "Tables list options side by side so you can compare them.",
        ],
    )


def section_glossary(doc: Document) -> None:
    add_heading(doc, "Glossary", level=1)
    doc.add_paragraph(
        "Words and short forms used in this report, in plain language."
    )

    add_table(
        doc,
        ["Term", "What it means"],
        [
            ["AAC", "Augmentative and Alternative Communication. Tools that help people "
                    "who cannot speak easily. Includes speech-generating devices."],
            ["AdSense", "Google's advertising program. YouTube videos earn small amounts "
                        "of money each time ads play during them."],
            ["Affiliate link", "A special link that pays you a small commission when "
                               "someone buys the product through it."],
            ["AT", "Assistive Technology. Any device or software that helps people with "
                   "disabilities. MouthPad is AT."],
            ["ATIA", "Assistive Technology Industry Association. The largest yearly "
                     "conference for AT in the United States. Held in Orlando in January."],
            ["ATP", "Assistive Technology Professional. A certified clinician who knows "
                    "AT well enough to recommend devices for individual users."],
            ["B2B", "Business to Business. Selling to companies or institutions instead "
                    "of to individual people. snowflower's main income will come from B2B."],
            ["BoFu / ToFu", "Bottom of Funnel / Top of Funnel. BoFu = articles for people "
                            "ready to buy. ToFu = articles for curious newcomers. BoFu makes "
                            "more money."],
            ["CSUN", "California State University Northridge — also the name of a major "
                     "AT conference held in Anaheim every March."],
            ["Cron", "A way to schedule programs to run automatically. 'Cron-publishable' "
                     "means we can post automatically without a human pressing publish."],
            ["DME", "Durable Medical Equipment. A US insurance category that includes "
                    "many AT devices and decides what insurance will pay for."],
            ["Engagement", "Likes, comments, shares, replies. Platform algorithms reward "
                           "engagement with more reach."],
            ["FDA", "US Food and Drug Administration. Regulates medical devices."],
            ["FTC", "US Federal Trade Commission. Enforces rules on advertising disclosure."],
            ["IEP", "Individualized Education Program. The legal document that describes "
                    "what a US public school will provide for a student with disabilities."],
            ["KEAD", "한국장애인고용공단. Korean Employment Agency for Persons with "
                    "Disabilities. Provides funding for AT devices for workers."],
            ["KFTC", "Korea Fair Trade Commission. Korean equivalent of the US FTC."],
            ["KNAT", "국립재활원 중앙보조기기센터. Korean center that provides AT devices "
                    "through welfare programs."],
            ["LMN", "Letter of Medical Necessity. A document a doctor or therapist writes "
                    "to convince an insurance company to pay for a device."],
            ["MFDS", "식약처. Korean Ministry of Food and Drug Safety. Regulates medical "
                    "devices in Korea."],
            ["MouthPad", "augmental.tech's product. A small device worn on the roof of "
                         "the mouth that lets a person move a computer cursor with their tongue."],
            ["OAuth", "A standard way for one program to safely log in to another. We "
                      "use OAuth so snowflower's engine can post to YouTube and LinkedIn "
                      "on our behalf."],
            ["OT", "Occupational Therapist. A medical professional who helps people do "
                   "daily-life activities. OTs are often the people who recommend "
                   "MouthPad-style devices."],
            ["RESNA", "Rehabilitation Engineering and Assistive Technology Society of "
                      "North America. The group that certifies ATPs."],
            ["SCI", "Spinal Cord Injury. A common reason someone might need hands-free "
                    "input technology."],
            ["SLP", "Speech-Language Pathologist. A clinician who works with people who "
                    "have communication difficulties. SLPs recommend AAC devices."],
            ["VA", "US Department of Veterans Affairs. The largest single buyer of AT "
                   "devices in the United States."],
        ],
        col_widths=[1.3, 5.5],
    )


def section_executive_summary(doc: Document) -> None:
    add_heading(doc, "1 · Executive Summary", level=1)
    doc.add_paragraph(
        "snowflower is a YouTube channel and email newsletter about devices that let "
        "people use computers without their hands. We will write reviews, comparisons, "
        "and how-to guides. The audience is twofold: people with disabilities who want "
        "to learn about these devices, and the doctors and therapists who recommend "
        "them."
    )
    doc.add_paragraph(
        "We are also building a Python program — the snowflower engine — that turns "
        "one source video or article into many posts across many platforms (YouTube, "
        "LinkedIn, Bluesky, newsletter, and more). One person can run the whole "
        "channel using this engine, with help from one paid co-host who has a "
        "disability."
    )
    doc.add_paragraph(
        "After 90 days of regular publishing, we will offer this audience to a "
        "company called augmental.tech. They make MouthPad — a device that lets a "
        "person move a computer cursor by moving their tongue. We will offer them "
        "three options: pay us as a sponsor, hire snowflower's editor, or partner "
        "with us on content. If they say no, we offer the same deal to their "
        "competitors (Tobii, Quha, Glassouse)."
    )

    add_callout(
        doc,
        "What we corrected from the first plan",
        "The original idea was to use AI to make lots of content, hide ads inside it, "
        "and use psychological tricks to grow fast. We rejected this for three reasons. "
        "First, hidden ads are illegal in the US, Korea, and Europe. Second, the "
        "platforms now punish AI-generated content that has no human voice — that "
        "shortcut closed in 2024 and 2025. Third, augmental sells accessibility "
        "hardware, so the audience is doctors and disabled users, not general tech "
        "fans. The corrected plan is open, transparent editorial covering the whole "
        "category fairly.",
    )

    add_heading(doc, "Key decisions", level=2)
    add_bullets(
        doc,
        [
            "Brand name: snowflower (눈꽃 in Korean). All accounts use this name only.",
            "We will hire a paid co-host who has a disability. They appear on camera "
            "and have the final say on how they are shown. Without this, we cannot "
            "build credibility with the disability community.",
            "We focus on motor disability and hands-free input. We do not try to "
            "cover all disabilities. Narrow focus wins faster.",
            "Korea is a future market. We will publish some Korean content for "
            "search visibility but will not push hard until augmental has a Korean "
            "distributor.",
            "Budget: $50 to $100 per month, plus $200+ per episode for the co-host.",
        ],
    )

    add_plain_words(
        doc,
        "snowflower is a small editorial channel that helps disabled people choose the "
        "right hands-free device, and helps the therapists who guide them. We give "
        "honest reviews, cover competitors fairly, and use a software tool we built "
        "to publish to many places at once. After 3 months we ask augmental.tech to "
        "pay us, hire us, or partner with us — using the audience as proof.",
    )


def section_big_picture(doc: Document) -> None:
    add_heading(doc, "2 · The Big Picture", level=1)

    add_heading(doc, "What is the problem?", level=2)
    doc.add_paragraph(
        "A person with a spinal cord injury cannot use a regular keyboard or mouse. "
        "They need a different way to control a computer. Several devices exist — "
        "MouthPad lets them use their tongue, Quha lets them use head movements, "
        "Tobii uses their eyes, Glassouse uses cheek muscles. Each device works "
        "well for some people and poorly for others. There is no single right answer."
    )
    doc.add_paragraph(
        "But there is no clear, trustworthy place online to compare these devices. "
        "Most reviews are written by people who tried the device for one hour. "
        "There is almost no information on how to get insurance to pay for the "
        "device, or how to talk to your therapist about it. The companies that make "
        "the devices are too small to advertise widely. The people who need the "
        "devices spend months looking for answers."
    )

    add_heading(doc, "What is snowflower's answer?", level=2)
    doc.add_paragraph(
        "We become the trusted reviewer for this category. We test every device "
        "honestly. We work with disabled testers who actually use the devices for "
        "weeks. We publish how to get each device paid for in each country, in each "
        "insurance system. We make the comparisons that no one else makes."
    )
    doc.add_paragraph(
        "Our content helps the user decide. Helps the therapist decide what to "
        "recommend. Helps the company sell more units to the right people. Everyone "
        "wins — except channels that produce shallow content."
    )

    add_heading(doc, "Why will this work in 2026?", level=2)
    add_bullets(
        doc,
        [
            "Big tech reviewers do not cover assistive devices in depth — too small a "
            "market for them. So the field is open.",
            "AI tools let one person do work that used to need a team. We can write, "
            "edit videos, generate Korean voiceovers, and publish to many platforms "
            "with software help. The cost is low.",
            "The companies we cover (augmental, Tobii, Quha) need credible third-party "
            "voices because their own marketing reads as biased. We provide that voice.",
            "Search engines and AI assistants are beginning to cite specialized "
            "publications more than general ones. Being narrow and deep is now an "
            "advantage.",
        ],
    )


def section_background(doc: Document) -> None:
    add_heading(doc, "3 · Background — Company, Product, Market", level=1)

    add_heading(doc, "3.1 augmental.tech and MouthPad", level=2)
    doc.add_paragraph(
        "augmental.tech is a small company that came out of MIT Media Lab. Their "
        "product is the MouthPad — the first hands-free computer trackpad worn "
        "inside the mouth. It is shaped like a thin dental retainer. The user moves "
        "their tongue across it, and a cursor moves on their phone or computer "
        "screen. Quick sips of air are clicks. The device weighs about 7.5 grams "
        "and is about 1 millimeter thick. It connects to almost any modern device "
        "by Bluetooth and lasts 5 hours per charge."
    )
    doc.add_paragraph(
        "Stories about augmental have appeared in The Verge, MIT News, BBC, NBC "
        "News, Forbes, TechCrunch, and Wired. The most-shared story is about "
        "Keely Horch, a college student at the University of Maryland, who uses "
        "MouthPad to study, work, and stay in touch with family. The story is "
        "warm, but augmental still struggles to reach the medical professionals "
        "who actually decide what device a patient gets."
    )

    add_heading(doc, "3.2 The competing devices", level=2)
    add_table(
        doc,
        ["Type", "Examples", "Who it works best for"],
        [
            ["Worn in the mouth", "MouthPad",
             "People with very limited body movement (high SCI, ALS, locked-in)"],
            ["Worn on the head", "Quha Zono 2, SmartNav 4, headband sensors",
             "People who can move their head reliably (most SCI, many CP, MS)"],
            ["Cheek button on glasses", "Glassouse V2",
             "People who can move their head and clench cheek muscles"],
            ["Eye tracking", "Tobii Eye Tracker 5, Tobii PCEye, Eyegaze Edge",
             "People who can move their eyes well; great for ALS and locked-in"],
            ["Built into your phone", "Apple AssistiveTouch, Voice Control, Switch Control; Windows Eye Control; Android Switch Access",
             "Anyone with a smartphone — totally free, but limited"],
            ["Voice control", "Apple Voice Control, Google Voice Access, Talon Voice",
             "People who can speak; works well for typing and navigation"],
            ["Sip and puff", "Origin Instruments Jouse, AbleNet QuadJoy",
             "Older technology, still used; reliable but slower"],
        ],
        col_widths=[1.6, 3.0, 2.4],
    )

    add_heading(doc, "3.3 How these devices get paid for", level=2)
    doc.add_paragraph(
        "Almost no one pays $5,000 cash for a MouthPad. The money comes from one "
        "of these places, and every path needs a doctor or therapist's signature."
    )
    add_table(
        doc,
        ["Path", "Who decides", "Typical wait"],
        [
            ["US private insurance", "Insurance company (after doctor + SLP/OT eval)",
             "1–4 months, often denied first"],
            ["US Medicare / Medicaid", "Federal/state agency (DME approval process)",
             "2–6 months"],
            ["US Veterans Affairs (VA)", "VA case manager + VA OT or prosthetist",
             "1–3 months"],
            ["US K-12 schools (IEP)", "School district AT specialist",
             "Decided each year in the IEP meeting"],
            ["State Vocational Rehab", "VR counselor + assessment",
             "1–4 months"],
            ["Korea — KEAD funding", "한국장애인고용공단 evaluator",
             "Up to 15M KRW per worker (20M for severe disability)"],
            ["Korea — KNAT welfare", "국립재활원 중앙보조기기센터",
             "Varies by device class"],
            ["Korea — 산재보험", "Workers' comp insurance",
             "For occupational injury cases"],
        ],
        col_widths=[2.0, 3.0, 2.0],
    )

    add_plain_words(
        doc,
        "MouthPad is a small computer device worn in the mouth. It competes with "
        "head-controlled, eye-controlled, and software-only options. People rarely "
        "buy these themselves — they go through insurance, the VA, schools, or "
        "Korean government programs. A doctor or therapist's signature is needed "
        "every single time. So our marketing must reach those signers, not just "
        "the patients.",
    )


def section_strategic_premise(doc: Document) -> None:
    add_heading(doc, "4 · The Three Mistakes We Avoided", level=1)
    doc.add_paragraph(
        "When we first discussed this project, three ideas seemed appealing. All "
        "three were wrong. Here is what each was, why it broke, and what we did "
        "instead."
    )

    add_heading(doc, "4.1 Mistake one: hidden ads and psychological tricks", level=2)
    doc.add_paragraph(
        "The original plan included 'dark psychology to grab attention' and "
        "'secret ads' inside the content. This is illegal in the US (FTC rules), "
        "in Korea (KFTC rules), and in Europe (DSA rules). Even worse, hidden ads "
        "are now caught easily by AI tools that look at content patterns. When "
        "caught, the company's reputation is damaged for years."
    )
    doc.add_paragraph(
        "Replacement: every snowflower post says clearly whether we received money, "
        "a free sample, or nothing. We have a public disclosure document with "
        "templates for each case, in English and Korean."
    )

    add_heading(doc, "4.2 Mistake two: automate everything to grow fast", level=2)
    doc.add_paragraph(
        "The original plan was to use AI to make many posts on every platform "
        "every day. This worked in 2022 and 2023. It does not work now. In 2024, "
        "YouTube banned thousands of channels that posted 'mass-produced AI "
        "content.' Instagram lowered the reach of obviously automated posts. X "
        "(Twitter) made the API expensive. TikTok wrote a new rule that bans apps "
        "that copy content from other platforms."
    )
    doc.add_paragraph(
        "Replacement: pick a few platforms, post less but better. Tier 1 = YouTube "
        "+ LinkedIn + Bluesky + newsletter. Use the engine to save time on the "
        "boring parts, but keep humans (us + the disabled co-host) in charge of "
        "what is actually said."
    )

    add_heading(doc, "4.3 Mistake three: wrong audience", level=2)
    doc.add_paragraph(
        "The original niche was 'AI and digital twin.' But augmental.tech does not "
        "sell AI software — they sell a tongue-controlled trackpad for people who "
        "cannot use their hands. The audiences are completely different. Posting "
        "about AI to AI fans would not bring augmental any customers."
    )
    doc.add_paragraph(
        "Replacement: the real niche is 'hands-free and assistive input "
        "technology.' Our audience is people with motor disabilities, their "
        "families, and the therapists who help them."
    )

    add_callout(
        doc,
        "Why this matters",
        "Many marketing projects fail because the team starts building before "
        "checking the basic facts. We checked first, found three big mistakes, and "
        "saved months of wasted work. The cleaner version of the plan is also more "
        "honest — which makes it easier to keep going for the long time it takes "
        "to build a real audience.",
        SUCCESS_HEX,
    )


def section_brand(doc: Document) -> None:
    add_heading(doc, "5 · Brand and Strategy", level=1)

    add_heading(doc, "5.1 Two things we are building at the same time", level=2)
    add_table(
        doc,
        ["What", "Description", "Why it matters"],
        [
            ["The engine",
             "A Python software tool that takes one piece of source content and "
             "publishes it to many platforms in the right format for each one. "
             "The code is in a public GitHub repository.",
             "If our channel ever gets sold or licensed, the engine is the asset. "
             "Augmental could even take it in-house and run it themselves."],
            ["The channel",
             "snowflower itself — the actual YouTube channel, LinkedIn page, "
             "Bluesky account, and newsletter that publishes content using the "
             "engine.",
             "The channel is the proof. It shows the engine works in real life "
             "with real readers, and that is what makes the engine worth buying."],
        ],
        col_widths=[1.2, 3.5, 2.3],
    )

    add_heading(doc, "5.2 Editorial rules we will not break", level=2)
    add_bullets(
        doc,
        [
            "Hire a paid disabled co-host before the first episode airs. They have "
            "the final word on how they appear in the video. We pay at least $200 "
            "per episode to start, and increase as the channel grows.",
            "Always cover competitors. If we review MouthPad, we mention or compare "
            "to Tobii, Quha, and Glassouse. This sounds like it would help "
            "competitors — but it actually makes us more trustworthy, which makes "
            "augmental want us more, not less.",
            "Use the language each person prefers. Some communities prefer "
            "'autistic person.' Others prefer 'person with cerebral palsy.' We "
            "ask each individual subject what they want, and we follow their lead.",
            "Never use words that the disability community has rejected. No "
            "'inspirational,' 'overcoming,' 'wheelchair-bound,' 'high-functioning,' "
            "or 'low-functioning.' No before/after framing that treats disability "
            "as a problem to fix.",
            "If we cannot prove a medical claim, we attribute it to whoever made "
            "the claim. We do not invent health benefits.",
        ],
    )

    add_heading(doc, "5.3 The 90-day pitch arc", level=2)
    doc.add_paragraph(
        "We run snowflower for 90 days. We publish weekly. We attend the CSUN AT "
        "conference in March. We collect data on who is watching: how many "
        "occupational therapists, how many speech-language pathologists, how many "
        "school AT specialists, how many VA staff. We build a picture of our "
        "audience that augmental cannot get any other way."
    )
    doc.add_paragraph(
        "Then in week 12 we send augmental a short, clear pitch: 'Here is the "
        "audience we built. Here is what they watched. Here is your competitors' "
        "presence in our content. Pay us as a sponsor, hire our editor, or "
        "partner with us on content.' If they say no, the same pitch goes to "
        "Tobii, Quha, and Glassouse — they all face the same problem augmental "
        "does."
    )


def section_platform(doc: Document, matrix_path: Path) -> None:
    add_heading(doc, "6 · Where We Will Publish", level=1)
    doc.add_paragraph(
        "We looked at every platform that could host our content. Eleven made the "
        "shortlist. They are sorted into three tiers based on how easy it is to "
        "publish to them automatically, and how well they match our audience."
    )

    add_figure(
        doc, matrix_path,
        "Figure 6.1 — Each dot is a platform. The top right is the best place to "
        "publish: easy to automate AND a good audience match. Bottom left is "
        "skip-or-defer.",
        width_inches=6.5,
    )

    add_heading(doc, "6.1 Tier 1 — start here", level=2)
    doc.add_paragraph("These four platforms have free or low-cost APIs and a strong audience match.")
    add_table(
        doc,
        ["Platform", "What we publish", "Why it matters"],
        [
            ["YouTube",
             "Long videos (10–20 min) with full reviews, plus short videos (under "
             "60 seconds) for quick demos.",
             "Most people search YouTube for 'how does X device work' before they "
             "buy. Our videos must be the first ones they find."],
            ["LinkedIn (Company Page)",
             "Articles and posts aimed at occupational therapists, school AT "
             "specialists, VA staff, and insurance reviewers.",
             "These are the people who decide whether a patient gets a device. "
             "They use LinkedIn for professional reading."],
            ["Bluesky",
             "Daily short posts, conference live-coverage, conversation with "
             "researchers.",
             "Most accessibility researchers and AT-twitter migrated to Bluesky "
             "in 2024–2025. The audience is small but exactly right."],
            ["Beehiiv (newsletter)",
             "Weekly long-form essay covering one topic in depth, plus a monthly "
             "'what's new' summary.",
             "Email goes directly to the reader. Algorithms can change; an email "
             "list is ours forever."],
        ],
        col_widths=[1.5, 2.8, 2.5],
    )

    add_heading(doc, "6.2 Tier 2 — add later, after Tier 1 works", level=2)
    add_table(
        doc,
        ["Platform", "Wait reason", "Wait time"],
        [
            ["Threads", "Comes free with Instagram", "After Instagram is set up"],
            ["Instagram Reels", "Meta requires an 'App Review' for our app",
             "1 to 4 weeks"],
            ["TikTok", "TikTok requires an audit before allowing automatic posting; "
             "until approved, all our posts are private",
             "5 to 10 business days"],
            ["X (Twitter)", "Posting through software now requires paying for the "
             "API ($200+ per month minimum for new users)",
             "Instant once we pay; defer until we have revenue"],
        ],
        col_widths=[1.5, 4.0, 1.4],
    )

    add_heading(doc, "6.3 Tier 3 — manual or Korea-specific", level=2)
    add_table(
        doc,
        ["Platform", "Why it is here"],
        [
            ["Reddit (r/disability, r/als, r/SCI, etc.)",
             "Reddit punishes accounts that post promotional links too often. We "
             "must post by hand, no more than once per week per subreddit, after "
             "an account has built up reputation."],
            ["Hacker News",
             "Manual submission only; once per quarter, when we publish a "
             "particularly important piece."],
            ["Naver Blog (Korean)",
             "In 2025 Naver started banning blogs that posted automated content. "
             "We will write Korean posts by hand, no more than 1–2 per month."],
            ["Stibee (Korean newsletter)",
             "API works, but we wait until we have a Korean email list."],
            ["KakaoStory, Tistory, Brunch, Disquiet, ResearchGate",
             "All have either dead APIs, no APIs, or no audience for our niche. We "
             "skip them entirely."],
        ],
        col_widths=[2.4, 4.4],
    )

    add_plain_words(
        doc,
        "We focus our energy on YouTube, LinkedIn, Bluesky, and our newsletter. "
        "These four are easy to publish to and reach the right people. Other "
        "platforms either cost too much, take too long to set up, or have wrong "
        "audiences. We can always add more platforms later — it is much cheaper "
        "to start narrow and expand than to start wide and fail.",
    )


def section_engine(doc: Document) -> None:
    add_heading(doc, "7 · The Software Engine", level=1)
    doc.add_paragraph(
        "snowflower is not just content. It is also a piece of software we built. "
        "Most of this report's value comes from this software, because it is what "
        "lets one person run a multi-platform channel. The software is open source "
        "and lives at github.com/ksk5429/snowflower."
    )

    add_heading(doc, "7.1 What the engine does, in plain words", level=2)
    doc.add_paragraph(
        "You give the engine one episode (a video file, an article, or both). It:"
    )
    add_numbered(
        doc,
        [
            "Reads the source content from a simple file (an episode YAML file).",
            "Generates captions automatically using the Whisper speech-to-text model.",
            "Cuts the long video into short vertical clips for TikTok and Reels.",
            "Generates a thumbnail picture using fal.ai (an AI image generator).",
            "Writes the post in the right format for each platform — short for "
            "Bluesky, longer for LinkedIn, even longer for the newsletter.",
            "Posts to each platform using its official API.",
            "Saves a record of what was posted, so we can measure results later.",
        ],
    )

    add_heading(doc, "7.2 What is in the code", level=2)
    add_table(
        doc,
        ["File group", "Purpose"],
        [
            ["snowflower.py", "The main program. Run this from the command line."],
            ["models.py, base_connector.py", "Defines what an Episode and a Post "
                                              "look like in the code."],
            ["connector_*.py (11 files)", "One file per platform — each knows how "
                                          "to talk to that platform's API."],
            ["transformer_*.py (5 files)", "Files that turn one type of content into "
                                            "another (text into captions, video into shorts, etc.)."],
            ["auth_*.py (3 files)", "Helper scripts that handle the safe login "
                                     "process for YouTube, LinkedIn, and Bluesky."],
            ["test_connectors.py", "Automated tests. Run them to make sure nothing "
                                    "is broken before publishing."],
            ["methodology.md, pitch_deck.md, etc.", "Strategy and reference "
                                                     "documents (including this report)."],
        ],
        col_widths=[2.5, 4.4],
    )

    add_callout(
        doc,
        "Verified working today",
        "All 33 automated tests pass. We can show what would be posted to all 11 "
        "platforms without actually posting (called 'dry-run mode'). Korean text "
        "renders properly on Windows. Running the engine on the first episode "
        "produces output for all 11 platforms in less than 1 second.",
        SUCCESS_HEX,
    )


def section_sample_content(doc: Document) -> None:
    add_heading(doc, "8 · What snowflower's Content Will Actually Look Like", level=1)
    doc.add_paragraph(
        "Strategy documents are useful but abstract. This section shows what one "
        "real snowflower episode would look like across each Tier 1 platform. "
        "These are sample drafts — the actual content will be tested and refined "
        "with real disabled co-hosts."
    )

    add_heading(doc, "8.1 Source episode: 'MouthPad vs Quha — which one for ALS?'", level=2)
    doc.add_paragraph(
        "This is a comparison episode aimed at people in the early stages of ALS, "
        "their families, and the speech-language pathologists who often help them "
        "make this decision."
    )

    add_heading(doc, "8.2 What the YouTube version looks like", level=3)
    add_callout(
        doc,
        "Title: 'MouthPad vs Quha Zono — which works longer as ALS progresses?'",
        "Length: 12 minutes. Format: side-by-side demo with disabled co-host using "
        "both devices for the same tasks (writing an email, browsing the web, "
        "scrolling Instagram). Shows speed numbers from our methodology page. "
        "Discusses how each device performs at month 1 vs month 6 vs month 12 of "
        "ALS progression. Ends with: 'Talk to your SLP about which one fits your "
        "current movement abilities.'",
        ACCENT_HEX,
    )

    add_heading(doc, "8.3 What the LinkedIn version looks like", level=3)
    add_callout(
        doc,
        "Title: 'Why ALS device decisions are easier when you start with the "
        "long view'",
        "A 600-word article aimed at SLPs. Argues that the device-recommendation "
        "conversation should anticipate disease progression, not just current "
        "ability. Cites our methodology page's fatigue measurements. Includes a "
        "downloadable PDF table showing typical performance trajectories. Ends "
        "with a quiet call to read the YouTube video for full demonstrations.",
        ACCENT_HEX,
    )

    add_heading(doc, "8.4 What the Bluesky version looks like", level=3)
    add_callout(
        doc,
        "Bluesky thread (4 posts of 280 characters each)",
        "Post 1: 'New deep-dive — MouthPad vs Quha for ALS, focusing on what "
        "happens 6–12 months into the disease, not just the unboxing demo.' "
        "Post 2: One key data point (a fatigue chart). Post 3: A quote from the "
        "co-host. Post 4: Link to YouTube + link to newsletter signup. Self-labels "
        "#a11y and #disabilitytech. Replies to every comment within 60 minutes.",
        ACCENT_HEX,
    )

    add_heading(doc, "8.5 What the newsletter version looks like", level=3)
    add_callout(
        doc,
        "Subject line: 'The ALS device choice everyone gets wrong'",
        "Long-form essay (1200 words) that goes deeper than the YouTube video. "
        "Covers: the underdiscussed problem of fatigue accumulation, the role of "
        "the SLP in long-term planning, three real user stories (with permission), "
        "specific funding paths in the US and Korea, and a closing 'next-step' "
        "checklist for readers and a separate one for clinicians.",
        ACCENT_HEX,
    )

    add_heading(doc, "8.6 What the disclosure footer looks like (every post)", level=2)
    add_callout(
        doc,
        "Disclosure (English):",
        "snowflower is independent editorial. We are not affiliated with augmental, "
        "Tobii, Quha, or any device covered. We received no payment, samples, or "
        "pre-publication review. snowflower.tech/about",
        ACCENT_HEX,
    )
    add_callout(
        doc,
        "Disclosure (Korean):",
        "snowflower(눈꽃)은 독립 에디토리얼 채널입니다. 본 콘텐츠에 등장하는 어떤 "
        "브랜드와도 협찬·후원·샘플 제공·사전 검토 관계가 없습니다.",
        ACCENT_HEX,
    )

    add_plain_words(
        doc,
        "One source video becomes 4–6 different posts. Each post is shaped for "
        "the platform it lives on. The disclosure is the same on every one. The "
        "engine handles the boring formatting work; the human writes the script "
        "and chooses the angle.",
    )


def section_research(doc: Document, journey_path: Path) -> None:
    add_heading(doc, "9 · What Professional Marketing Teaches Us", level=1)
    doc.add_paragraph(
        "Before writing snowflower's plan, we read the playbooks of four kinds "
        "of professional marketers: disability-focused content creators, "
        "independent tech publications, B2B medical-device marketers, and "
        "newsletter monetization experts. Here are the lessons that mattered "
        "most."
    )

    add_heading(doc, "9.1 The seven biggest lessons", level=2)
    add_table(
        doc,
        ["Lesson", "Plain-language explanation"],
        [
            ["Hire a disabled co-host before episode 1.",
             "Channels about disability that are run only by non-disabled people "
             "are not trusted. Fix this before launch, not after."],
            ["Pick a narrow niche, not a broad one.",
             "We cover hands-free input only. We do not cover deafness, blindness, "
             "or other disabilities. Narrow channels reach 10,000 readers faster "
             "than broad channels reach 100."],
            ["Publish a methodology page before reviews.",
             "Tell people exactly how we test. Use scientific instruments where "
             "possible (Fitts's law, NASA-TLX). Publish the raw data. This is "
             "what makes a tech reviewer trustworthy."],
            ["Write articles for buyers, not browsers.",
             "Articles like 'MouthPad vs Quha' make money. Articles like 'What is "
             "assistive technology?' don't. Same effort, ten times the return."],
            ["Email is the asset; everything else is acquisition.",
             "YouTube, LinkedIn, and Bluesky can change their rules tomorrow. An "
             "email list is ours. Use the platforms to get email subscribers."],
            ["Audience density matters more than audience size.",
             "2,500 readers who are all therapists is worth more than 200,000 "
             "random tech fans. Augmental will pay for the first kind, not the second."],
            ["Anchor the launch to a conference.",
             "CSUN AT 2026 in March is the biggest accessibility-tech event in the "
             "world. We will get press credentials and film the entire show floor."],
        ],
        col_widths=[2.4, 4.4],
    )

    add_heading(doc, "9.2 Who actually buys these devices, step by step", level=2)
    add_figure(
        doc, journey_path,
        "Figure 9.1 — How a person actually ends up with a MouthPad in the United "
        "States. The therapist signature is the bottleneck on every path.",
        width_inches=6.6,
    )

    add_heading(doc, "9.3 Where the buyers actually read", level=2)
    add_table(
        doc,
        ["Where", "Who reads it", "Can we publish there?"],
        [
            ["The ASHA Leader (US)",
             "Speech-language pathologists",
             "Yes — they accept guest editorial pieces"],
            ["OT Practice / AOTA (US)",
             "Occupational therapists",
             "Yes — they have a published rate card for sponsored content"],
            ["AT Today (UK)",
             "AT professionals in the UK",
             "Yes — display ads + advertorial"],
            ["Closing The Gap (US)",
             "School AT specialists",
             "Yes — formal advertising program"],
            ["에이블뉴스 (Korea)",
             "Korean disability community + sector workers",
             "Yes — formal 광고/후원 form"],
            ["함께걸음, 더인디고 (Korea)",
             "Korean disability advocacy",
             "Yes — by inquiry"],
        ],
        col_widths=[1.8, 2.7, 2.5],
    )

    add_plain_words(
        doc,
        "Professional marketers in this niche have already figured out what works: "
        "narrow focus, deep methodology, articles aimed at buyers, email above "
        "everything else, and trade publications that reach the people who write "
        "the prescriptions. We are following these lessons exactly.",
    )


def section_methodology(doc: Document) -> None:
    add_heading(doc, "10 · How We Test Devices", level=1)
    doc.add_paragraph(
        "Most accessibility-tech reviews say things like 'felt good' or 'easy to "
        "use.' That is not enough for a therapist deciding whether to recommend "
        "a $5,000 device that insurance will scrutinize. snowflower publishes a "
        "complete test methodology so our results are repeatable and disputable."
    )

    add_heading(doc, "10.1 The six things we measure", level=2)
    add_table(
        doc,
        ["#", "What we measure", "How we measure it (in plain words)"],
        [
            ["1", "Performance",
             "Speed (how many bits per second the user can move and click) and "
             "typing rate (words per minute, error-corrected). We use a standard "
             "test from a 2012 international standard."],
            ["2", "Setup time",
             "How long from opening the box to the first successful click? "
             "Stopwatch timed."],
            ["3", "Sustainability",
             "Does the speed stay steady, or does the user get tired? We test "
             "after 1, 15, and 30 minutes. We also use two questionnaires that "
             "measure mental and physical effort."],
            ["4", "Compatibility",
             "Does it work on macOS, iOS, Windows, Android, and ChromeOS? With "
             "common apps like Office, Zoom, Slack, Safari?"],
            ["5", "Funding",
             "How much will the user actually pay out of pocket under each "
             "insurance / VA / school / Korean funding path?"],
            ["6", "Practical daily use",
             "Battery life, charging time, weight, comfort during long sessions, "
             "comfort using it in public."],
        ],
        col_widths=[0.3, 1.6, 4.9],
    )

    add_heading(doc, "10.2 The most important rule", level=2)
    doc.add_paragraph(
        "A non-disabled person testing a MouthPad will get certain numbers. A "
        "person with a C2 spinal cord injury will get different numbers. Both are "
        "real. But only the second matters to a person trying to decide whether "
        "to buy one. Every snowflower review will lead with numbers from the "
        "disabled co-host. Numbers from non-disabled testers go in an appendix."
    )

    add_callout(
        doc,
        "We publish the raw data",
        "Every review includes the actual measurement files: raw click logs, "
        "questionnaire responses, the exact software and firmware versions used. "
        "Anyone — a vendor, a researcher, a competitor — can re-run our tests and "
        "challenge our numbers. We publish methodology disputes and our responses "
        "publicly on GitHub.",
        SUCCESS_HEX,
    )


def section_audience_density(doc: Document) -> None:
    add_heading(doc, "11 · The Audience-Density Idea", level=1)

    add_heading(doc, "11.1 Why size is not what matters", level=2)
    doc.add_paragraph(
        "Most YouTube channels measure success in subscribers. snowflower will "
        "measure success in 'audience density.' This is the share of our audience "
        "that consists of people who actually buy or recommend assistive devices: "
        "occupational therapists, speech-language pathologists, school AT "
        "specialists, VA prosthetists, and the like."
    )
    doc.add_paragraph(
        "Why? Because augmental.tech does not need general awareness. The Verge, "
        "BBC, and Forbes already covered them. What augmental needs is direct "
        "reach to the small group of people who decide whether a patient gets a "
        "MouthPad. That group is at most 200,000 people in the US — but each one "
        "of them is worth more than 1,000 random viewers."
    )

    add_heading(doc, "11.2 Example calculation", level=2)
    add_table(
        doc,
        ["Channel A", "Channel B", "Who augmental will pay more"],
        [
            ["200,000 general tech YouTube subscribers, mostly hobbyists",
             "2,500 newsletter readers, of whom 800 are verified therapists",
             "Channel B — by 5–10×, easily"],
            ["Earns about $200/month in YouTube ads",
             "Earns $3,000–10,000/month from one anchor sponsor like augmental",
             "Channel B"],
            ["Algorithm risk if YouTube changes rules",
             "Owns the email list, no algorithm risk",
             "Channel B"],
        ],
        col_widths=[2.4, 2.4, 2.0],
    )

    add_heading(doc, "11.3 Who exactly we count", level=2)
    add_bullets(
        doc,
        [
            "Speech-Language Pathologists (SLPs) — the main gatekeeper for AAC "
            "device funding.",
            "Occupational Therapists (OTs) — the main gatekeeper for access "
            "devices like MouthPad.",
            "Assistive Technology Professionals (ATPs) — RESNA-certified "
            "specialists who recommend devices.",
            "School District AT Specialists — decide what AT a public school "
            "buys for students.",
            "VA prosthetists, OTs, and case managers — the largest single buyer "
            "in the US.",
            "Rehab engineers and educators at SCI rehab centers (Shepherd, Craig, "
            "Kessler, Magee, Spaulding).",
            "Korean equivalents: 작업치료사 (OT), 언어재활사 (SLP), 보조공학사 "
            "(KR ATP), KEAD evaluators, 재활의학과 전문의.",
        ],
    )

    add_plain_words(
        doc,
        "Augmental does not need millions of viewers. They need the right few "
        "thousand. snowflower's whole strategy is to be the channel that has those "
        "right few thousand, with proof of who they are. That is what we will "
        "show in the pitch.",
    )


def section_90day(doc: Document, timeline_path: Path) -> None:
    add_heading(doc, "12 · The 90-Day Plan, Week by Week", level=1)

    add_figure(
        doc, timeline_path,
        "Figure 12.1 — The full 90-day plan. Conference at week 9 is the centerpiece. "
        "Pitch to augmental in week 12 with three months of data behind it.",
        width_inches=6.8,
    )

    add_heading(doc, "12.1 Week-by-week tasks", level=2)
    add_table(
        doc,
        ["Week", "What we do"],
        [
            ["W1",
             "Buy snowflower.tech domain. Create gmail and Google Voice number. "
             "Open Bluesky account. Open YouTube Brand Account and set up Google "
             "Cloud project for the API. Identify a candidate co-host and start "
             "the conversation. Draft methodology page."],
            ["W2",
             "Set up Beehiiv newsletter account. Outline the first 4 episodes. "
             "Apply for press credentials at CSUN AT 2026 (deadline approaches "
             "fast — apply early). Draft pitch letters to ASHA Leader and OT "
             "Practice for guest articles."],
            ["W3",
             "Record episode 1 (Apple AssistiveTouch primer). Publish methodology "
             "page so it is live before any review. Post first 5 introduction "
             "posts on Bluesky."],
            ["W4",
             "Episode 1 publishes on YouTube. First newsletter goes out. LinkedIn "
             "Company Page launches with first 3 posts."],
            ["W5",
             "Episode 2: 'What MouthPad actually does — explained simply.' "
             "Newsletter #2. 2 LinkedIn posts. Daily Bluesky."],
            ["W6",
             "Episode 3: 'Quha Zono 2 review — the head mouse alternative.' Same "
             "weekly rhythm."],
            ["W7",
             "Episode 4: 'Tobii PCEye review — eyes vs tongue vs head.' Continue."],
            ["W8",
             "Episode 5: 'Funding paths for SGDs in 2026 — US and Korea.'"],
            ["W9",
             "CSUN AT 2026 conference (March 9–13, Anaheim). Film 8–10 vendor "
             "booths, interview 5 clinicians, try to get an on-camera moment with "
             "augmental's founders if they attend."],
            ["W10",
             "Edit conference footage into the start of a 6-episode series. "
             "Newsletter feature: 'What I learned at CSUN 2026.'"],
            ["W11",
             "Continue editing CSUN series. Reach out to interview subjects for "
             "follow-up content."],
            ["W12",
             "Send pitch to augmental: data deck (audience composition, OT/SLP "
             "percentage, MouthPad-coverage views), 90-day metrics, three "
             "engagement options. If they respond: schedule call. If they don't: "
             "give them a week, then send the same pitch to Tobii or Quha."],
        ],
        col_widths=[0.6, 6.2],
    )

    add_heading(doc, "12.2 What success looks like at day 90", level=2)
    add_callout(
        doc,
        "Realistic targets",
        "1,500–3,000 YouTube subscribers. 800–2,000 LinkedIn Company Page "
        "followers. 500–1,500 Bluesky followers. 800–2,500 newsletter "
        "subscribers, of whom 100–400 are verified clinicians. 8–12 published "
        "episodes including 6 from CSUN. One pitch sent to augmental, with data "
        "behind it.",
        SUCCESS_HEX,
    )


def section_budget(doc: Document, budget_path: Path, revenue_path: Path) -> None:
    add_heading(doc, "13 · Money — Budget and Revenue", level=1)

    add_heading(doc, "13.1 What we spend each month", level=2)
    doc.add_paragraph(
        "Total monthly software cost is between $30 and $90, depending on which "
        "extras we use. The biggest single recurring cost is our newsletter "
        "platform once we exceed 2,500 subscribers. Outside of software, we pay "
        "the disabled co-host at least $200 per episode."
    )

    add_figure(
        doc, budget_path,
        "Figure 13.1 — Monthly budget by category, all under the $50–100 cap.",
        width_inches=6.5,
    )

    add_heading(doc, "13.2 What an example month looks like", level=2)
    add_table(
        doc,
        ["Item", "Frequency", "Cost"],
        [
            ["Beehiiv newsletter (after 2,500 subs)", "Monthly", "$39"],
            ["fal.ai image generation", "Per thumbnail (~$0.50)", "~$18"],
            ["ElevenLabs Korean voice", "Per minute of audio", "~$22"],
            ["Cloud / GitHub Actions", "Monthly", "$6"],
            ["X API Basic", "Monthly", "$17"],
            ["Captioning / misc", "Monthly", "$8"],
            ["Subtotal software", "", "~$110"],
            ["Co-host fee", "Per episode (4/month)", "$200 × 4 = $800"],
            ["Total month example", "", "~$910"],
        ],
        col_widths=[3.0, 2.5, 1.5],
    )
    doc.add_paragraph(
        "Software cost can be reduced to about $30/month by skipping ElevenLabs "
        "(if Korean is not yet active) and X API (if we delay the X automation). "
        "The co-host cost is non-negotiable and rises with audience size."
    )

    add_heading(doc, "13.3 How revenue should grow", level=2)
    add_figure(
        doc, revenue_path,
        "Figure 13.2 — Estimated revenue mix month by month. The big jump in "
        "month 8 represents the augmental anchor sponsorship if the pitch works. "
        "If they pass, we replace this with another sponsor.",
        width_inches=6.6,
    )

    add_heading(doc, "13.4 What we will charge once we start charging", level=2)
    add_bullets(
        doc,
        [
            "Newsletter sponsor placement: $200–400 per email at 1,000–2,500 "
            "subscribers. Higher when our subscribers are mostly clinicians.",
            "LinkedIn sponsored post: $500–1,500 each at 5,000 followers.",
            "YouTube sponsor mention in a video: $500–2,000 at 5,000–25,000 subs.",
            "Anchor sponsor package (target: month 6 or later): $3,000–10,000 "
            "per month for 6 months minimum, with rates tied to measured website "
            "traffic delivered.",
            "We do not price by views (CPM). We price by audience density. "
            "Sponsors pay more for an audience full of buyers than for a bigger "
            "audience full of browsers.",
        ],
    )

    add_plain_words(
        doc,
        "We can run the engine for $30–90/month in software. Adding the human "
        "co-host pushes the real monthly cost to roughly $1,000. Revenue starts "
        "small (YouTube ads, affiliate links) and grows toward newsletter "
        "sponsors and one anchor sponsor by month 6–8. The whole plan only works "
        "if the channel reaches the right people, not the most people.",
    )


def section_risk(doc: Document) -> None:
    add_heading(doc, "14 · What Could Go Wrong (and What We Do About It)", level=1)
    doc.add_paragraph(
        "Every plan has risks. The honest thing is to name them up front. Here are "
        "the nine biggest risks for snowflower, ranked by likelihood, with our "
        "plan for each."
    )

    add_table(
        doc,
        ["Risk", "How likely", "What we do"],
        [
            ["The disability community feels we are an outsider channel "
             "exploiting their stories",
             "Medium-high",
             "Hire a paid disabled co-host before episode 1. Give them final cut "
             "on every scene featuring them. Publish an open complaints process. "
             "Take public corrections seriously."],
            ["augmental rejects the pitch at day 90",
             "Medium",
             "The engine, the audience, and the methodology are all valuable "
             "assets even without augmental. The same pitch goes to Tobii, Quha, "
             "or Glassouse next. We also build other revenue (Patreon, "
             "newsletter sponsors)."],
            ["YouTube wrongly demonetizes our videos as 'sensitive content'",
             "Medium",
             "Plan AdSense as no more than 30% of revenue. Patreon and direct "
             "sponsorships are primary income."],
            ["Naver Blog suspends us if we post too automatically",
             "Medium-high if we automate",
             "We do not automate Naver. The connector refuses to publish — by "
             "design. Korean Naver posts are written and posted by hand only."],
            ["LinkedIn does not approve our developer access",
             "Medium",
             "Even without approval, the basic posting works. Full Company Page "
             "API is a nice-to-have, not a must-have."],
            ["TikTok audit fails or takes months",
             "Medium",
             "We can post manually without API. We delay TikTok until YouTube and "
             "Bluesky have proven the content works."],
            ["augmental complains about how we use their trademark",
             "Low",
             "We never use 'MouthPad' or 'augmental' as our account name or logo. "
             "We always disclose our independent status. The DISCLOSURE.md file "
             "documents the rules."],
            ["The single editor (you) burns out at month 4–6",
             "High if we publish too aggressively",
             "Pre-make 4 episodes before launch. Outsource video editing once "
             "revenue covers it ($50–150 per video). Skip Tier 2 platforms until "
             "Tier 1 is steady."],
            ["Korean market entry blocked because MouthPad lacks 식약처 (MFDS) "
             "clearance",
             "External — depends on augmental",
             "Korea is build-awareness mode only. We do not run paid Korean "
             "campaigns until augmental has a Korean distributor. We continue "
             "publishing Korean content to be ready when they do."],
        ],
        col_widths=[2.4, 1.4, 3.0],
    )

    add_callout(
        doc,
        "The biggest single risk is our own pace",
        "More content channels die from burnout than from algorithm changes or "
        "competition. The plan is intentionally modest in cadence: 1 deep "
        "episode per week, 1 newsletter, 2 LinkedIn posts, daily Bluesky. We "
        "pre-make 4 episodes before launch so we do not start behind.",
        DANGER_HEX,
    )


def section_compliance(doc: Document) -> None:
    add_heading(doc, "15 · Legal Rules We Follow", level=1)

    add_heading(doc, "15.1 Disclosure rules — three regions", level=2)
    add_table(
        doc,
        ["Region", "Law / regulator", "What it requires"],
        [
            ["United States", "FTC 16 CFR § 255",
             "If a brand pays you, gives you a free sample, or is your employer, "
             "you must disclose this clearly inside every post — not just in the "
             "bio."],
            ["Korea", "표시·광고법 (KFTC)",
             "Same idea — paid content must be marked clearly. Naver and Kakao "
             "also lower the reach of undisclosed ads, separately from the law."],
            ["European Union", "Digital Services Act Art. 26",
             "Sponsored content must be transparent. Applies to anyone whose "
             "content is shown to EU users."],
        ],
        col_widths=[1.5, 2.0, 3.4],
    )

    add_heading(doc, "15.2 Medical claim rules", level=2)
    doc.add_paragraph(
        "MouthPad is a body-contact electronic device. Claims about how well it "
        "works for medical conditions are regulated by the US FDA. Our rule is "
        "simple: we never invent or rephrase medical claims. If we want to say "
        "'MouthPad helps X people do Y,' we must be quoting augmental's published "
        "source directly, with the citation. We never claim medical outcomes as "
        "our own findings."
    )

    add_heading(doc, "15.3 Brand-only identity — one exception", level=2)
    add_table(
        doc,
        ["Platform", "Brand-only OK?", "Why"],
        [
            ["YouTube (Brand Account), TikTok Business, Instagram Creator, "
             "Threads, Bluesky, X, Reddit, Substack/Beehiiv",
             "Yes",
             "Standard for editorial brands. The platform allows a brand-only "
             "account name."],
            ["LinkedIn personal profile",
             "No",
             "LinkedIn requires real names for personal profiles. Fake-name "
             "personal accounts are terminated under their Terms of Service "
             "section 8.2."],
            ["LinkedIn Company Page",
             "Yes",
             "Brand-only is exactly what a Company Page is for. We use Company "
             "Page only and skip personal posting on LinkedIn."],
        ],
        col_widths=[3.0, 1.4, 2.5],
    )

    add_callout(
        doc,
        "Why these rules are also good business",
        "Following disclosure law is not just about avoiding fines. It is what "
        "makes the audience trust us. And the trust is what makes augmental "
        "willing to pay us. The companies that get caught hiding ads lose far "
        "more value than they would have made.",
        SUCCESS_HEX,
    )


def section_faq(doc: Document) -> None:
    add_heading(doc, "16 · Frequently Asked Questions", level=1)

    qs = [
        ("Is snowflower a real company or just a project?",
         "Right now snowflower is a personal project — a publishing brand and a "
         "software engine. If the augmental pitch leads to anything serious, we "
         "may register a real company (LLC in the US, or 1인 사업자 in Korea). "
         "For now, the engine and content live on GitHub under one person's "
         "account."),
        ("Why hire a disabled co-host? Isn't that expensive?",
         "It is the single most important spending decision. Without a disabled "
         "co-host of record, channels covering disability are correctly criticized "
         "as outsiders. The trust ceiling without one is low. With one, we can "
         "approach Apple, Microsoft, and brand sponsors at the highest tier. The "
         "$200 per episode pays for itself within a few months."),
        ("Why not just sponsor MouthPad directly without all this work?",
         "We do not have augmental's permission to act as their official channel. "
         "Doing so would be impersonation and trademark misuse. Building an "
         "independent channel that they choose to engage with later is the only "
         "legally clean version of this."),
        ("Why Beehiiv instead of Substack?",
         "Substack takes 10% of all your subscription revenue forever. Beehiiv "
         "takes 0%. Beehiiv also has a built-in advertising network that fills "
         "automatically. For our scale, Beehiiv pays for itself by 1,500 "
         "subscribers and is ours forever after that."),
        ("Why include Korean content if we are mostly US-focused?",
         "Two reasons. First, almost no Western AT publication covers the Korean "
         "funding system. That is a unique angle that makes us valuable to "
         "augmental as they think about international expansion. Second, the "
         "channel owner is Korean and can write fluent Korean, so the marginal "
         "cost is low."),
        ("Why not use AI to write everything?",
         "We use AI for the boring parts: captions, image thumbnails, language "
         "translation. We do not use AI for the writing voice, the choice of "
         "what to cover, or the medical claims. AI-only content is now actively "
         "down-ranked on YouTube and LinkedIn since 2024. More importantly, our "
         "credibility comes from the human voice."),
        ("What happens if augmental simply ignores the pitch?",
         "Three backup paths. First, the same pitch goes to Tobii Dynavox, "
         "Quha, or Glassouse — they have the same problem augmental does. "
         "Second, we monetize through Patreon-style memberships from clinics "
         "and schools. Third, we sell sponsored guest articles to ASHA Leader "
         "and OT Practice. The plan does not depend on augmental saying yes."),
        ("Is 90 days enough time?",
         "It is enough time to send a credible pitch with real data — that is "
         "the goal. It is not enough to become the dominant channel; that takes "
         "12–18 months. The 90-day moment is the first real check on whether the "
         "approach works. If augmental engages, we accelerate. If not, we "
         "continue building and approach competitors."),
        ("What is the minimum to start?",
         "Buy the domain ($15), make a Gmail account (free), get a Google Voice "
         "number (free), open a Bluesky account (free), generate the app "
         "password (free). Then publish the first text-only post on Bluesky. "
         "Total cost: $15. Total time: about 30 minutes."),
        ("Who is the target reader of this report?",
         "Three audiences. First, the channel owner (you), to keep the strategy "
         "clear over the long execution. Second, anyone we hire — the disabled "
         "co-host, an editor, a Korean translator. Third, augmental and other "
         "potential sponsors, who see this report and understand exactly what "
         "they are buying into."),
    ]
    for q, a in qs:
        add_heading(doc, q, level=3)
        doc.add_paragraph(a)


def section_conclusion(doc: Document) -> None:
    add_heading(doc, "17 · Conclusion and Next Actions", level=1)
    doc.add_paragraph(
        "snowflower exists to fill a real gap. There is no other publication "
        "that combines deep technical reviews of hands-free input devices, "
        "clinician-readable methodology, fair coverage of competitors, and "
        "Korean-language coverage of the Korean funding system. The engine and "
        "the strategy are designed to compound across all four advantages, with "
        "augmental.tech as the first pitch target and clear backup paths if they "
        "decline."
    )

    add_callout(
        doc,
        "The smallest first step",
        "Buy snowflower.tech for $15. Create snowflower.editorial@gmail.com. "
        "Open a free Bluesky account, generate an app password, paste it into "
        "the .env file. Run 'python auth_bluesky.py' to verify the connection. "
        "Then run 'python snowflower.py publish --episode ep001_episode.yaml "
        "--live --platforms bluesky' to publish the first real post. Total "
        "cost: $15. Total time: about 30 minutes.",
        SUCCESS_HEX,
    )

    add_heading(doc, "Next actions in order", level=2)
    add_numbered(
        doc,
        [
            "Buy snowflower.tech via Porkbun or Namecheap (~$15/year).",
            "Create snowflower.editorial@gmail.com plus Google Voice number.",
            "Create a Bluesky account, mint an app password, fill in .env, run "
            "python auth_bluesky.py.",
            "Publish episode 1 to Bluesky in live mode — proves the pipeline "
            "works against a real API.",
            "Open YouTube Brand Account + Google Cloud project + OAuth (about "
            "10 minutes). Run python auth_youtube.py.",
            "Open LinkedIn Company Page (skip personal profile per LinkedIn ToS).",
            "Identify and contract a paid disabled co-host before episode 2.",
            "Apply for CSUN AT 2026 press credentials (deadline approaches "
            "fast — apply early in week 2).",
            "Pitch ASHA Leader and OT Practice for guest editorial slots in "
            "week 2.",
            "Pre-make episodes 1–4 before public launch — survives the early "
            "burnout window where most niche channels die.",
        ],
    )


def section_appendix(doc: Document) -> None:
    add_heading(doc, "Appendix A · What is in the snowflower repository", level=1)
    doc.add_paragraph(
        "Repository: github.com/ksk5429/snowflower (public). Layout is "
        "intentionally flat — all files are at the top level, no subfolders for "
        "the source code. Run scripts directly using python <filename>.py."
    )
    add_table(
        doc,
        ["Document", "Purpose"],
        [
            ["README.md", "Project overview, quick-start commands"],
            ["DISCLOSURE.md", "FTC + KFTC + EU disclosure templates (4 kinds × "
                              "EN/KR)"],
            ["methodology.md", "Test methodology framework (the credibility "
                                "anchor)"],
            ["accounts_setup.md", "Per-platform account-creation work-along "
                                   "checklist"],
            ["pitch_deck.md", "augmental.tech pitch (audience-density-first)"],
            ["research_findings.md", "Synthesized 4-stream marketing research "
                                      "playbook"],
            ["template_titles.md", "Pain Point SEO patterns + LinkedIn hook "
                                    "templates + ban list"],
            ["template_disclosure.md", "Copy-paste disclosure footers"],
            ["ep001_episode.yaml", "First episode source — Apple AssistiveTouch "
                                    "primer"],
            ["snowflower_strategy_report.docx", "This document — generated by "
                                                  "generate_report.py"],
        ],
        col_widths=[2.4, 4.4],
    )

    add_heading(doc, "Appendix B · References", level=1)
    add_bullets(
        doc,
        [
            "augmental.tech and MouthPad (MIT News, June 2024)",
            "FTC 16 CFR § 255 — Endorsement Guides",
            "Korea Fair Trade Commission 추천·보증 등에 관한 광고 심사지침",
            "ISO 9241-411:2012 — Evaluation methods for the design of physical "
            "input devices",
            "NASA-TLX (Hart and Staveland, 1988) — workload assessment",
            "Borg CR10 scale — perceived exertion",
            "ATIA 2026 (January, Orlando), CSUN AT 2026 (March, Anaheim), RESNA "
            "2026 (March, Long Beach), Closing The Gap 2026 (October, Minneapolis)",
            "ASHA Leader, AOTA OT Practice, AT Today UK, 에이블뉴스, 함께걸음, "
            "더인디고",
            "RTINGS methodology pages (rtings.com/labs)",
            "Wirecutter founding-era review structure (Brian Lam, 2011–2016)",
            "Stratechery (Ben Thompson) — paid newsletter analytical model",
            "Justin Welsh / Lara Acosta / Sahil Bloom — 2026 LinkedIn B2B "
            "playbooks",
            "Grow & Convert — Pain Point SEO methodology",
            "HubSpot — topic-cluster architecture",
            "MrBeast leaked memo (2023) — first 15 seconds + thumbnail rules",
        ],
    )


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------


def main() -> int:
    print("Generating charts...")
    matrix_path = chart_platform_matrix()
    timeline_path = chart_90_day_timeline()
    budget_path = chart_budget_breakdown()
    journey_path = chart_buyer_journey()
    revenue_path = chart_revenue_mix()
    print(f"  - {matrix_path}")
    print(f"  - {timeline_path}")
    print(f"  - {budget_path}")
    print(f"  - {journey_path}")
    print(f"  - {revenue_path}")

    print("Building document...")
    doc = Document()
    set_core_properties(doc)
    configure_styles(doc)

    # Cover (section 0)
    build_cover(doc)

    # Body section
    new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    new_section.top_margin = Inches(0.85)
    new_section.bottom_margin = Inches(0.85)
    new_section.left_margin = Inches(0.85)
    new_section.right_margin = Inches(0.85)
    new_section.header.is_linked_to_previous = False
    new_section.footer.is_linked_to_previous = False

    # TOC
    toc_para = doc.add_paragraph()
    toc_run = toc_para.add_run("Table of Contents")
    toc_run.bold = True
    toc_run.font.size = Pt(20)
    toc_run.font.color.rgb = PRIMARY_RGB
    doc.add_paragraph()
    toc_field_para = doc.add_paragraph()
    _add_toc_field(toc_field_para)
    p_note = doc.add_paragraph(
        "(In Microsoft Word, right-click anywhere in the table above and select "
        "'Update Field' to fill it with current page numbers.)"
    )
    p_note.runs[0].italic = True
    p_note.runs[0].font.size = Pt(9)
    p_note.runs[0].font.color.rgb = MUTED_RGB

    # Front matter
    section_readers_guide(doc)
    section_glossary(doc)

    # Body — 17 sections
    section_executive_summary(doc)
    section_big_picture(doc)
    section_background(doc)
    section_strategic_premise(doc)
    section_brand(doc)
    section_platform(doc, matrix_path)
    section_engine(doc)
    section_sample_content(doc)
    section_research(doc, journey_path)
    section_methodology(doc)
    section_audience_density(doc)
    section_90day(doc, timeline_path)
    section_budget(doc, budget_path, revenue_path)
    section_risk(doc)
    section_compliance(doc)
    section_faq(doc)
    section_conclusion(doc)
    section_appendix(doc)

    configure_header_footer(doc)

    print(f"Saving {REPORT_PATH}...")
    doc.save(REPORT_PATH)

    size_kb = REPORT_PATH.stat().st_size / 1024
    print(f"OK wrote {REPORT_PATH} ({size_kb:.1f} KB)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
